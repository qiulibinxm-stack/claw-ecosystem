#!/usr/bin/env python3
# 生成酒店调研和评估报告的简化Python脚本

import json
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def load_data():
    """加载数据"""
    with open('hotel_deep_data.json', 'r', encoding='utf-8') as f:
        hotel_data = json.load(f)
    with open('hotel_estimations.json', 'r', encoding='utf-8') as f:
        estimations = json.load(f)
    return hotel_data, estimations

def format_number(num):
    """格式化数字"""
    return f"{num:,}"

def create_deep_research_report():
    """创建深度调研报告"""
    print("正在创建深度调研报告...")
    
    hotel_data, estimations = load_data()
    
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # 封面
    title = doc.add_heading('酒店深度调研报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('苏福比画廊酒店 & 北海银滩吉海度假酒店')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('万能虾运营分析团队')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'生成时间: {datetime.datetime.now().strftime("%Y年%m月%d日")}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 一、执行摘要
    doc.add_heading('一、执行摘要', 1)
    doc.add_paragraph('本报告基于纯自动化数据采集和分析，对苏福比画廊酒店和北海银滩吉海度假酒店进行深度调研，为运营决策提供数据支持。')
    
    est1 = estimations["苏福比画廊酒店"]["estimated_metrics"]
    est2 = estimations["北海银滩吉海度假酒店"]["estimated_metrics"]
    
    doc.add_heading('1.1 关键发现', 2)
    doc.add_paragraph('• 两个酒店均位于北海银滩旅游核心区，地理位置优越')
    doc.add_paragraph('• 竞争环境激烈：500米内均有15个竞争对手')
    doc.add_paragraph('• 周边配套完善：餐饮、购物、交通设施齐全')
    doc.add_paragraph(f'• 经营指标预估：苏福比画廊酒店年营收{format_number(est1["annual_revenue_est"])}元，北海银滩吉海度假酒店年营收{format_number(est2["annual_revenue_est"])}元')
    
    # 二、酒店基本信息
    doc.add_heading('二、酒店基本信息', 1)
    
    hotel1 = hotel_data["苏福比画廊酒店"]
    hotel2 = hotel_data["北海银滩吉海度假酒店"]
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '指标'
    hdr_cells[1].text = '苏福比画廊酒店'
    hdr_cells[2].text = '北海银滩吉海度假酒店'
    
    # 填充数据
    rows = table.rows
    rows[1].cells[0].text = '地址'
    rows[1].cells[1].text = hotel1["basic_info"]["address"]
    rows[1].cells[2].text = hotel2["basic_info"]["address"]
    
    rows[2].cells[0].text = '电话'
    rows[2].cells[1].text = hotel1["basic_info"]["tel"]
    rows[2].cells[2].text = hotel2["basic_info"]["tel"]
    
    rows[3].cells[0].text = '评分'
    rows[3].cells[1].text = hotel1["basic_info"]["rating"] + '分'
    rows[3].cells[2].text = hotel2["basic_info"]["rating"] + '分'
    
    rows[4].cells[0].text = '坐标'
    rows[4].cells[1].text = hotel1["basic_info"]["location"]
    rows[4].cells[2].text = hotel2["basic_info"]["location"]
    
    # 三、经营指标预估
    doc.add_heading('三、经营指标预估', 1)
    
    table2 = doc.add_table(rows=8, cols=3)
    table2.style = 'Table Grid'
    
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = '经营指标'
    hdr_cells2[1].text = '苏福比画廊酒店'
    hdr_cells2[2].text = '北海银滩吉海度假酒店'
    
    rows2 = table2.rows
    rows2[1].cells[0].text = '预估房间数'
    rows2[1].cells[1].text = str(est1["room_count_est"]) + '间'
    rows2[1].cells[2].text = str(est2["room_count_est"]) + '间'
    
    rows2[2].cells[0].text = '平均房价'
    rows2[2].cells[1].text = str(est1["avg_daily_rate_est"]) + '元'
    rows2[2].cells[2].text = str(est2["avg_daily_rate_est"]) + '元'
    
    rows2[3].cells[0].text = '预估入住率'
    rows2[3].cells[1].text = f'{est1["occupancy_rate_est"]*100:.1f}%'
    rows2[3].cells[2].text = f'{est2["occupancy_rate_est"]*100:.1f}%'
    
    rows2[4].cells[0].text = '预估年营收'
    rows2[4].cells[1].text = format_number(est1["annual_revenue_est"]) + '元'
    rows2[4].cells[2].text = format_number(est2["annual_revenue_est"]) + '元'
    
    rows2[5].cells[0].text = '预估年利润'
    rows2[5].cells[1].text = format_number(est1["annual_profit_est"]) + '元'
    rows2[5].cells[2].text = format_number(est2["annual_profit_est"]) + '元'
    
    rows2[6].cells[0].text = '预估利润率'
    rows2[6].cells[1].text = f'{est1["profit_margin_est"]}%'
    rows2[6].cells[2].text = f'{est2["profit_margin_est"]}%'
    
    rows2[7].cells[0].text = '预估投资回报率'
    rows2[7].cells[1].text = f'{est1["roi_est"]}%'
    rows2[7].cells[2].text = f'{est2["roi_est"]}%'
    
    # 四、优势与风险
    doc.add_heading('四、优势与风险分析', 1)
    
    doc.add_heading('4.1 苏福比画廊酒店', 2)
    doc.add_paragraph('优势：艺术主题差异化明显，避免同质化竞争；评分4.6分显示客户满意度良好；周边景点丰富（4个），有利于吸引游客；利润率较高（50%），盈利能力强。')
    doc.add_paragraph('风险：艺术主题受众可能有限；需要艺术内容运营，复杂度较高；竞争激烈（竞争系数0.425）。')
    
    doc.add_heading('4.2 北海银滩吉海度假酒店', 2)
    doc.add_paragraph('优势：评分极高（4.8分），客户体验优秀；度假酒店定位清晰；投资回报率较高（24%）；环境配套完善。')
    doc.add_paragraph('风险：度假酒店竞争更激烈；投资额较高；季节性明显，淡旺季差异大。')
    
    # 五、运营建议
    doc.add_heading('五、运营建议', 1)
    
    doc.add_heading('5.1 苏福比画廊酒店运营建议', 2)
    doc.add_paragraph('• 深化艺术主题：定期举办艺术展览、工作坊')
    doc.add_paragraph('• 差异化定价：利用艺术主题实现溢价')
    doc.add_paragraph('• 跨界合作：与艺术机构、画廊合作')
    doc.add_paragraph('• 内容营销：打造艺术主题社交媒体内容')
    
    doc.add_heading('5.2 北海银滩吉海度假酒店运营建议', 2)
    doc.add_paragraph('• 强化度假体验：提供一站式度假服务')
    doc.add_paragraph('• 家庭客群开发：针对家庭度假需求设计产品')
    doc.add_paragraph('• 旺季策略优化：最大化旺季收益')
    doc.add_paragraph('• 会员体系建立：提高客户复购率')
    
    # 保存文档
    doc.save('酒店深度调研报告.docx')
    print("✓ 深度调研报告已生成：酒店深度调研报告.docx")
    
    return doc

def create_project_evaluation_report():
    """创建项目评估报告"""
    print("正在创建项目评估报告...")
    
    hotel_data, estimations = load_data()
    
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # 封面
    title = doc.add_heading('酒店项目运营评估报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('苏福比画廊酒店 & 北海银滩吉海度假酒店')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('万能虾投资决策团队')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'生成时间: {datetime.datetime.now().strftime("%Y年%m月%d日")}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 一、项目评估摘要
    doc.add_heading('一、项目评估摘要', 1)
    
    est1 = estimations["苏福比画廊酒店"]["estimated_metrics"]
    est2 = estimations["北海银滩吉海度假酒店"]["estimated_metrics"]
    
    doc.add_paragraph('基于深度调研数据，对两个酒店项目的运营价值进行综合评估，为投资决策提供依据。')
    
    doc.add_heading('1.1 核心评估结论', 2)
    doc.add_paragraph(f'• 苏福比画廊酒店：预估年营收{format_number(est1["annual_revenue_est"])}元，年利润{format_number(est1["annual_profit_est"])}元，投资回报率{est1["roi_est"]}%')
    doc.add_paragraph(f'• 北海银滩吉海度假酒店：预估年营收{format_number(est2["annual_revenue_est"])}元，年利润{format_number(est2["annual_profit_est"])}元，投资回报率{est2["roi_est"]}%')
    doc.add_paragraph('• 两个项目均具备运营价值，但需考虑竞争环境和季节性因素')
    
    # 二、投资价值对比
    doc.add_heading('二、投资价值对比', 1)
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '评估维度'
    hdr_cells[1].text = '苏福比画廊酒店'
    hdr_cells[2].text = '北海银滩吉海度假酒店'
    
    rows = table.rows
    rows[1].cells[0].text = '营收潜力'
    rows[1].cells[1].text = format_number(est1["annual_revenue_est"]) + '元'
    rows[1].cells[2].text = format_number(est2["annual_revenue_est"]) + '元'
    
    rows[2].cells[0].text = '利润水平'
    rows[2].cells[1].text = format_number(est1["annual_profit_est"]) + '元'
    rows[2].cells[2].text = format_number(est2["annual_profit_est"]) + '元'
    
    rows[3].cells[0].text = '投资回报率'
    rows[3].cells[1].text = f'{est1["roi_est"]}%'
    rows[3].cells[2].text = f'{est2["roi_est"]}%'
    
    rows[4].cells[0].text = '风险水平'
    rows[4].cells[1].text = '中'
    rows[4].cells[2].text = '中'
    
    rows[5].cells[0].text = '差异化程度'
    rows[5].cells[1].text = '高（艺术主题）'
    rows[5].cells[2].text = '中（度假主题）'
    
    # 三、风险评估
    doc.add_heading('三、风险评估', 1)
    
    doc.add_heading('3.1 市场风险', 2)
    doc.add_paragraph('• 竞争风险：500米内均有15个竞争对手，竞争激烈')
    doc.add_paragraph('• 季节性风险：北海旅游淡旺季差异大，淡季收入可能大幅下降')
    doc.add_paragraph('• 同质化风险：周边多为别墅型度假酒店，产品同质化严重')
    
    doc.add_heading('3.2 运营风险', 2)
    doc.add_paragraph('• 管理风险：需要专业的酒店管理团队')
    doc.add_paragraph('• 成本风险：人工、物料成本可能上涨')
    doc.add_paragraph('• 服务质量风险：需要保持高评分以维持竞争力')
    
    # 四、投资建议
    doc.add_heading('四、投资建议', 1)
    
    doc.add_heading('4.1 投资优先级', 2)
    if est1["roi_est"] > est2["roi_est"]:
        doc.add_paragraph('• 优先推荐：苏福比画廊酒店（投资回报率更高）')
        doc.add_paragraph('• 次选：北海银滩吉海度假酒店')
    else:
        doc.add_paragraph('• 优先推荐：北海银滩吉海度假酒店（投资回报率更高）')
        doc.add_paragraph('• 次选：苏福比画廊酒店')
    
    doc.add_heading('4.2 投资策略', 2)
    doc.add_paragraph('• 小规模试点：先尝试3-6个月合作，验证模型')
    doc.add_paragraph('• 业绩对赌：设计业绩对赌条款，降低风险')
    doc.add_paragraph('• 分阶段投资：根据业绩表现分阶段追加投资')
    doc.add_paragraph('• 退出机制：明确合作退出条件，保护投资')
    
    # 五、下一步行动
    doc.add_heading('五、下一步行动建议', 1)
    
    doc.add_paragraph('1. 实地考察：安排1-2天实地调研，验证数据准确性')
    doc.add_paragraph('2. 经营者访谈：了解实际经营情况和合作意愿')
    doc.add_paragraph('3. 财务数据验证：获取实际经营数据进行验证')
    doc.add_paragraph('4. 合作方案设计：设计具体的运营合作方案')
    doc.add_paragraph('5. 法律尽调：进行法律和合规性审查')
    
    # 保存文档
    doc.save('酒店项目评估报告.docx')
    print("✓ 项目评估报告已生成：酒店项目评估报告.docx")
    
    return doc

def main():
    """主函数"""
    print("=" * 50)
    print("开始生成酒店分析报告...")
    print(f"时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 50)
    
    try:
        # 创建深度调研报告
        create_deep_research_report()
        
        # 创建项目评估报告
        create_project_evaluation_report()
        
        print("\n" + "=" * 50)
        print("报告生成完成！")
        print("已生成以下文件：")
        print("1. 酒店深度调研报告.docx - 详细的市场调研分析")
        print("2. 酒店项目评估报告.docx - 投资价值和风险评估")
        print("=" * 50)
        
        # 列出生成的文件
        import os
        files = os.listdir('.')
        docx_files = [f for f in files if f.endswith('.docx')]
        print(f"\n当前目录中的Word文档：")
        for f in docx_files:
            size = os.path.getsize(f)
            print(f"  • {f} ({size:,} bytes)")
        
    except Exception as e:
        print(f"生成报告时出错: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()