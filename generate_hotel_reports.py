#!/usr/bin/env python3
# 生成酒店调研和评估报告的Python脚本

import json
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def load_data():
    """加载数据"""
    with open('hotel_deep_data.json', 'r', encoding='utf-8') as f:
        hotel_data = json.load(f)
    with open('hotel_estimations.json', 'r', encoding='utf-8') as f:
        estimations = json.load(f)
    return hotel_data, estimations

def create_deep_research_report(hotel_data, estimations):
    """创建深度调研报告"""
    print("创建深度调研报告...")
    
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # 封面
    title = doc.add_heading('酒店深度调研报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('苏福比画廊酒店 & 北海银滩吉海度假酒店')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('万能虾运营分析团队')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'生成时间: {datetime.datetime.now().strftime("%Y年%m月%d日")}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 目录
    doc.add_heading('目录', 1)
    doc.add_paragraph('一、执行摘要')
    doc.add_paragraph('二、酒店基本信息')
    doc.add_paragraph('三、竞争环境分析')
    doc.add_paragraph('四、周边环境分析')
    doc.add_paragraph('五、经营指标预估')
    doc.add_paragraph('六、优势与风险分析')
    doc.add_paragraph('七、运营建议')
    doc.add_paragraph('附录：数据来源与方法')
    
    doc.add_page_break()
    
    # 一、执行摘要
    doc.add_heading('一、执行摘要', 1)
    p = doc.add_paragraph('本报告基于纯自动化数据采集和分析，对苏福比画廊酒店和北海银滩吉海度假酒店进行深度调研，为运营决策提供数据支持。调研采用高德地图API、行业基准数据和环境分析模型，实现零人工输入的自动化分析。')
    
    doc.add_heading('1.1 关键发现', 2)
    doc.add_paragraph('• 两个酒店均位于北海银滩旅游核心区，地理位置优越')
    doc.add_paragraph('• 竞争环境激烈：500米内均有15个竞争对手')
    doc.add_paragraph('• 周边配套完善：餐饮、购物、交通设施齐全')
    doc.add_paragraph(f'• 经营指标预估：苏福比画廊酒店年营收{estimations["苏福比画廊酒店"]["estimated_metrics"]["annual_revenue_est"]:,}元，北海银滩吉海度假酒店年营收{estimations["北海银滩吉海度假酒店"]["estimated_metrics"]["annual_revenue_est"]:,}元')
    
    # 二、酒店基本信息
    doc.add_heading('二、酒店基本信息', 1)
    
    # 创建表格
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '指标'
    hdr_cells[1].text = '苏福比画廊酒店'
    hdr_cells[2].text = '北海银滩吉海度假酒店'
    
    # 填充数据
    hotel1 = hotel_data["苏福比画廊酒店"]
    hotel2 = hotel_data["北海银滩吉海度假酒店"]
    
    rows = table.rows
    rows[1].cells[0].text = '地址'
    rows[1].cells[1].text = hotel1["basic_info"]["address"]
    rows[1].cells[2].text = hotel2["basic_info"]["address"]
    
    rows[2].cells[0].text = '电话'
    rows[2].cells[1].text = hotel1["basic_info"]["tel"]
    rows[2].cells[2].text = hotel2["basic_info"]["tel"]
    
    rows[3].cells[0].text = '评分'
    rows[3].cells[1].text = hotel1["basic_info"]["rating"] + '分'
    rows[3].cells[2].text = hotel2["basic_info"]["rating"] + '分'
    
    rows[4].cells[0].text = '坐标'
    rows[4].cells[1].text = hotel1["basic_info"]["location"]
    rows[4].cells[2].text = hotel2["basic_info"]["location"]
    
    # 三、竞争环境分析
    doc.add_heading('三、竞争环境分析', 1)
    
    table2 = doc.add_table(rows=4, cols=3)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = '竞争指标'
    hdr_cells2[1].text = '苏福比画廊酒店'
    hdr_cells2[2].text = '北海银滩吉海度假酒店'
    
    rows2 = table2.rows
    rows2[1].cells[0].text = '500米内竞争对手'
    rows2[1].cells[1].text = str(len(hotel1["competitors"])) + '个'
    rows2[1].cells[2].text = str(len(hotel2["competitors"])) + '个'
    
    rows2[2].cells[0].text = '最近竞品距离'
    rows2[2].cells[1].text = hotel1["competitors"][0]["distance"] + '米' if hotel1["competitors"] else '未知'
    rows2[2].cells[2].text = hotel2["competitors"][0]["distance"] + '米' if hotel2["competitors"] else '未知'
    
    rows2[3].cells[0].text = '竞争系数'
    rows2[3].cells[1].text = str(estimations["苏福比画廊酒店"]["estimated_metrics"]["comp_factor"])
    rows2[3].cells[2].text = str(estimations["北海银滩吉海度假酒店"]["estimated_metrics"]["comp_factor"])
    
    # 四、周边环境分析
    doc.add_heading('四、周边环境分析', 1)
    
    table3 = doc.add_table(rows=6, cols=3)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells3 = table3.rows[0].cells
    hdr_cells3[0].text = '周边配套'
    hdr_cells3[1].text = '苏福比画廊酒店'
    hdr_cells3[2].text = '北海银滩吉海度假酒店'
    
    rows3 = table3.rows
    rows3[1].cells[0].text = '餐饮设施'
    rows3[1].cells[1].text = str(hotel1["surroundings"]["餐饮"]["count"]) + '家'
    rows3[1].cells[2].text = str(hotel2["surroundings"]["餐饮"]["count"]) + '家'
    
    rows3[2].cells[0].text = '购物设施'
    rows3[2].cells[1].text = str(hotel1["surroundings"]["购物"]["count"]) + '家'
    rows3[2].cells[2].text = str(hotel2["surroundings"]["购物"]["count"]) + '家'
    
    rows3[3].cells[0].text = '旅游景点'
    rows3[3].cells[1].text = str(hotel1["surroundings"]["景点"]["count"]) + '个'
    rows3[3].cells[2].text = str(hotel2["surroundings"]["景点"]["count"]) + '个'
    
    rows3[4].cells[0].text = '公交站点'
    rows3[4].cells[1].text = str(hotel1["traffic_info"]["bus_count"]) + '个'
    rows3[4].cells[2].text = str(hotel2["traffic_info"]["bus_count"]) + '个'
    
    rows3[5].cells[0].text = '环境综合得分'
    rows3[5].cells[1].text = str(estimations["苏福比画廊酒店"]["estimated_metrics"]["env_score"]) + '/1.4'
    rows3[5].cells[2].text = str(estimations["北海银滩吉海度假酒店"]["estimated_metrics"]["env_score"]) + '/1.4'
    
    # 五、经营指标预估
    doc.add_heading('五、经营指标预估', 1)
    
    est1 = estimations["苏福比画廊酒店"]["estimated_metrics"]
    est2 = estimations["北海银滩吉海度假酒店"]["estimated_metrics"]
    
    table4 = doc.add_table(rows=8, cols=3)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells4 = table4.rows[0].cells
    hdr_cells4[0].text = '经营指标'
    hdr_cells4[1].text = '苏福比画廊酒店'
    hdr_cells4[2].text = '北海银滩吉海度假酒店'
    
    rows4 = table4.rows
    rows4[1].cells[0].text = '预估房间数'
    rows4[1].cells[1].text = str(est1["room_count_est"]) + '间'
    rows4[1].cells[2].text = str(est2["room_count_est"]) + '间'
    
    rows4[2].cells[0].text = '平均房价'
    rows4[2].cells[1].text = str(est1["avg_daily_rate_est"]) + '元'
    rows4[2].cells[2].text = str(est2["avg_daily_rate_est"]) + '元'
    
    rows4[3].cells[0].text = '预估入住率'
    rows4[3].cells[1].text = f'{est1["occupancy_rate_est"]*100:.1f}%'
    rows4[3].cells[2].text = f'{est2["occupancy_rate_est"]*100:.1f}%'
    
    rows4[4].cells[0].text = '预估年营收'
    rows4[4].cells[1].text = f'{est1["annual_revenue_est"]:,}元'
    rows4[4].cells[2].text = f'{est2["annual_revenue_est"]:,}元'
    
    rows4[5].cells[0].text = '预估年利润'
    rows4[5].cells[1].text = f'{est1["annual_profit_est"]:,}元'
    rows4[5].cells[2].text = f'{est2["annual_profit_est"]:,}元'
    
    rows4[6].cells[0].text = '预估利润率'
    rows4[6].cells[1].text = f'{est1["profit_margin_est"]}%'
    rows4[6].cells[2].text = f'{est2["profit_margin_est"]}%'
    
    rows4[7].cells[0].text = '预估投资回报率'
    rows4[7].cells[1].text = f'{est1["roi_est"]}%'
    rows4[7].cells[2].text = f'{est2["roi_est"]}%'
    
    # 六、优势与风险分析
    doc.add_heading('六、优势与风险分析', 1)
    
    doc.add_heading('6.1 苏福比画廊酒店', 2)
    doc.add_paragraph('优势：艺术主题差异化明显，避免同质化竞争；评分4.6分显示客户满意度良好；周边景点丰富（4个），有利于吸引游客；利润率较高（50%），盈利能力强。')
    doc.add_paragraph('风险：艺术主题受众可能有限；需要艺术内容运营，复杂度较高；竞争激烈（竞争系数0.425）。')
    
    doc.add_heading('6.2 北海银滩吉海度假酒店', 2)
    doc.add_paragraph('优势：评分极高（4.8分），客户体验优秀；度假酒店定位清晰；投资回报率较高（24%）；环境配套完善。')
    doc.add_paragraph('风险：度假酒店竞争更激烈；投资额较高；季节性明显，淡旺季差异大。')
    
    doc.add_heading('6.3 共同风险', 2)
    doc.add_paragraph('• 竞争激烈：500米内均有15个竞争对手')
    doc.add_paragraph('• 季节性明显：北海旅游淡旺季差异大')
    doc.add_paragraph('• 同质化竞争：周边多为别墅型度假酒店')
    
    # 七、运营建议
    doc.add_heading('七、运营建议', 1)
    
    doc.add_heading('7.1 苏福比画廊酒店运营建议', 2)
    doc.add_paragraph('• 深化艺术主题：定期举办艺术展览、工作坊，打造艺术IP')
    doc.add_paragraph('• 差异化定价：利用艺术主题实现溢价，提高平均房价')
    doc.add_paragraph('• 跨界合作：与艺术机构、画廊合作，扩大影响力')
    doc.add_paragraph('• 内容营销：打造艺术主题社交媒体内容，吸引目标客群')
    
    doc.add_heading('7.2 北海银滩吉海度假酒店运营建议', 2)
    doc.add_paragraph('• 强化度假体验：提供一站式度假服务，增加附加价值')
    doc.add_paragraph('• 家庭客群开发：针对家庭度假需求设计产品和服务')
    doc.add_paragraph('• 旺季策略优化：最大化旺季收益，设计淡季促销活动')
    doc.add_paragraph('• 会员体系建立：提高客户复购率，建立忠诚度计划')
    
    # 附录
    doc.add_heading('附录：数据来源与方法', 1)
    
    doc.add_heading('数据来源', 2)
    doc.add_paragraph('• 高德地图API：实时地理位置、评分、电话等信息')
    doc.add_paragraph('• 行业基准数据：基于公开酒店行业报告')
    doc.add_paragraph('• 环境分析模型：自动化计算周边环境得分')
    doc.add_paragraph('• 竞争分析模型：基于竞争对手距离和数量')
    
    doc.add_heading('模型假设', 2)
    doc.add_paragraph('• 季节性假设：4-10月为旺季，11-3月为淡季')
    doc.add_paragraph('• 成本结构假设：基于行业平均成本率')
    doc.add_paragraph('• 投资估算假设：投资额=年营收×2')
    doc.add_paragraph('• 评分影响假设：高评分带来溢价和入住率提升')
    
    doc.add_heading('报告生成信息', 2)
    doc.add_paragraph(f'生成时间：{datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('生成方式：纯自动化，无需人工输入')
    doc.add_paragraph('模型版本：酒店经营预估模型 v1.0')
    doc.add_paragraph('生成团队：万能虾运营分析团队')
    
    # 保存文档
    doc.save('酒店深度调研报告.docx')
    print("深度调研报告已保存：酒店深度调研报告.docx")
    
    return doc

def create_project_evaluation_report(hotel_data, estimations):
    """创建项目评估报告"""
    print("创建项目评估报告...")
    
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # 封面
    title = doc.add_heading('酒店项目运营评估报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('苏福比画廊酒店 & 北海银滩吉海度假酒店')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('万能虾投资决策团队')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'生成时间: {datetime.datetime.now().strftime("%Y年%m月%d日")}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 一、项目评估摘要
    doc.add_heading('一、项目评估摘要', 1)
    
    est1 = estimations["苏福比画廊酒店"]["estimated_metrics"]
    est2 = estimations["北海银滩吉海度假酒店"]["estimated_metrics"]
    
    doc.add_paragraph('基于深度调研数据，对两个酒店项目的运营价值进行综合评估，为投资决策提供依据。')
    
    doc.add_heading('1.1 核心评估结论', 2)
    doc.add_paragraph(f'• 苏福比画廊酒店：预估年营收{est1["annual_revenue_est"]:,}元